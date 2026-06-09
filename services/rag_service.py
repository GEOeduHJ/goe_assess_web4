"""
지리 자동 채점 시스템의 RAG (검색 증강 생성) 서비스

업로드된 참고자료를 문서 청크로 변환하고 FAISS 기반 유사도 검색을 제공합니다.
RAGService는 세션 간 stale vector store 재사용을 피하기 위해 싱글턴으로 동작하지 않습니다.
"""

import os
import re
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional
import logging

from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document as LangChainDocument
from langchain_huggingface import HuggingFaceEmbeddings
from pypdf import PdfReader
from docx import Document
import tiktoken

from config import config


logger = logging.getLogger(__name__)


@dataclass
class RAGResult:
    """RAG 처리 결과"""

    success: bool
    content: List[str] = field(default_factory=list)
    error_message: str = ""


@dataclass
class ExtractedTextBlock:
    """문서에서 추출된 텍스트 블록과 메타데이터"""

    text: str
    metadata: Dict[str, object] = field(default_factory=dict)


class RAGService:
    """문서 처리 및 유사성 검색을 위한 RAG 서비스"""

    def __init__(self):
        self.embeddings = HuggingFaceEmbeddings(
            model_name=config.EMBEDDING_MODEL,
            encode_kwargs={
                "normalize_embeddings": config.FAISS_INDEX_TYPE == "IndexFlatIP",
                "batch_size": config.EMBEDDING_BATCH_SIZE,
            },
        )
        self.vector_store = None
        self.logger = logger
        self.tokenizer = tiktoken.get_encoding("cl100k_base")

    def process_documents(self, uploaded_files: List) -> bool:
        """업로드된 참고 문서를 처리하고 FAISS 벡터 저장소를 생성합니다."""
        try:
            documents: List[LangChainDocument] = []
            max_docs = getattr(config, "MAX_DOCS_PER_STUDENT", 0)
            files_to_process = uploaded_files[:max_docs] if max_docs and max_docs > 0 else uploaded_files

            for doc_index, file_obj in enumerate(files_to_process):
                try:
                    blocks = self._extract_document_blocks(file_obj)
                    chunk_count = 0
                    for block in blocks:
                        chunks = self._chunk_document(block.text)
                        for chunk_id, chunk in enumerate(chunks):
                            metadata = dict(block.metadata)
                            metadata.update(
                                {
                                    "source": file_obj.name,
                                    "doc_index": doc_index,
                                    "chunk_id": chunk_id,
                                }
                            )
                            documents.append(LangChainDocument(page_content=chunk, metadata=metadata))
                            chunk_count += 1
                            if chunk_count >= config.CHUNKS_PER_DOC_LIMIT:
                                break
                        if chunk_count >= config.CHUNKS_PER_DOC_LIMIT:
                            break
                except Exception as exc:
                    self.logger.warning("Skipping reference file during RAG processing: %s", exc)
                    continue

            if not documents:
                return False

            self.vector_store = FAISS.from_documents(documents, self.embeddings)
            return True
        except Exception as exc:
            self.logger.warning("RAG document processing failed: %s", exc)
            return False

    def search_relevant_content(self, query: str, k: Optional[int] = None) -> List[str]:
        """쿼리를 기반으로 관련 내용 검색"""
        try:
            if not self.vector_store or not query or not query.strip():
                return []

            k = k or config.TOP_K_RETRIEVAL
            docs = self.vector_store.similarity_search(query.strip(), k=k)
            results = []
            for doc in docs:
                source = doc.metadata.get("source", "unknown")
                page = doc.metadata.get("page")
                prefix = f"출처: {source}"
                if page:
                    prefix += f" p.{page}"
                results.append(f"{prefix}\n{doc.page_content}")
            return results
        except Exception as exc:
            self.logger.warning("RAG search failed: %s", exc)
            return []

    def process_documents_for_student(self, uploaded_files: List, student_answer: str) -> RAGResult:
        """문서를 처리하고 특정 학생 답안과 관련된 내용 검색"""
        try:
            if not self.vector_store:
                success = self.process_documents(uploaded_files)
                if not success:
                    return RAGResult(success=False, error_message="문서 처리 실패")

            relevant_content = self.search_relevant_content(student_answer)
            return RAGResult(success=True, content=relevant_content)
        except Exception as exc:
            return RAGResult(success=False, error_message=str(exc))

    def _extract_document_blocks(self, file_obj) -> List[ExtractedTextBlock]:
        """PDF 또는 DOCX 파일에서 텍스트 블록과 메타데이터를 추출합니다."""
        try:
            file_extension = Path(file_obj.name).suffix.lower()
            if hasattr(file_obj, "seek"):
                file_obj.seek(0)

            if file_extension == ".pdf":
                return self._extract_pdf_blocks(file_obj)
            if file_extension == ".docx":
                return self._extract_docx_blocks(file_obj)
            return []
        except Exception as exc:
            self.logger.warning("Failed to extract document blocks: %s", exc)
            return []

    def _extract_pdf_blocks(self, file_obj) -> List[ExtractedTextBlock]:
        """PDF 파일에서 페이지 단위 텍스트를 추출합니다."""
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            tmp_file.write(file_obj.read())
            tmp_file_path = tmp_file.name

        try:
            blocks: List[ExtractedTextBlock] = []
            with open(tmp_file_path, "rb") as pdf_file:
                pdf_reader = PdfReader(pdf_file)
                for page_index, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        blocks.append(
                            ExtractedTextBlock(
                                text=page_text,
                                metadata={"page": page_index, "section": f"page_{page_index}"},
                            )
                        )
            return blocks
        finally:
            os.unlink(tmp_file_path)

    def _extract_docx_blocks(self, file_obj) -> List[ExtractedTextBlock]:
        """DOCX 파일에서 문단과 표 내용을 추출합니다."""
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
            tmp_file.write(file_obj.read())
            tmp_file_path = tmp_file.name

        try:
            doc = Document(tmp_file_path)
            blocks: List[ExtractedTextBlock] = []

            paragraph_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraph_text.append(paragraph.text.strip())
            if paragraph_text:
                blocks.append(
                    ExtractedTextBlock(
                        text="\n\n".join(paragraph_text),
                        metadata={"section": "paragraphs"},
                    )
                )

            for table_index, table in enumerate(doc.tables, 1):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        rows.append(" | ".join(cells))
                if rows:
                    blocks.append(
                        ExtractedTextBlock(
                            text="\n".join(rows),
                            metadata={"section": f"table_{table_index}"},
                        )
                    )
            return blocks
        finally:
            os.unlink(tmp_file_path)

    def _chunk_document(
        self,
        content: str,
        chunk_tokens: Optional[int] = None,
        overlap_tokens: Optional[int] = None,
    ) -> List[str]:
        """문서 내용을 토큰 기반 겹침 청크로 분할합니다."""
        if not content or len(content.strip()) == 0:
            return []

        chunk_tokens = chunk_tokens or config.CHUNK_SIZE
        overlap_tokens = overlap_tokens if overlap_tokens is not None else config.CHUNK_OVERLAP
        if chunk_tokens <= 0:
            chunk_tokens = 300
        if overlap_tokens < 0:
            overlap_tokens = 0
        if overlap_tokens >= chunk_tokens:
            overlap_tokens = max(0, chunk_tokens // 5)

        content = clean_text(content)
        if not content:
            return []

        tokens = self.tokenizer.encode(content)
        if len(tokens) <= chunk_tokens:
            return [content]

        chunks = []
        start_idx = 0
        while start_idx < len(tokens):
            end_idx = min(start_idx + chunk_tokens, len(tokens))
            chunk_text = self.tokenizer.decode(tokens[start_idx:end_idx])
            if chunk_text.strip():
                chunks.append(chunk_text.strip())
            start_idx += chunk_tokens - overlap_tokens
        return chunks


def clean_text(text: str) -> str:
    """의미 있는 지리 기호를 보존하면서 공백을 정리합니다."""
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[^a-zA-Z0-9가-힣ㄱ-ㅎㅏ-ㅣ\s.!?,;:\-\(\)\[\]'\"°%℃㎜㎞㎡·→←↑↓/\\|+=<>~]", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def create_rag_service() -> RAGService:
    return RAGService()


def format_retrieved_content(content: List[str]) -> str:
    if not content:
        return ""

    formatted_chunks = []
    for i, chunk in enumerate(content, 1):
        chunk_text = chunk.strip()
        if chunk_text:
            formatted_chunks.append(f"참고자료 {i}:\n{chunk_text}")
    return "\n\n".join(formatted_chunks)
