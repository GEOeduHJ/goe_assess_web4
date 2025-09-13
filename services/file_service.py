"""
파일 처리 서비스 모듈
Excel 파일 업로드 및 형식 검증, PDF/DOCX 문서 내용 추출, 이미지 파일과 학생 이름 매칭 기능을 제공합니다.
"""

import os
import re
from typing import List, Dict, Optional, Tuple, Any
from pathlib import Path
import pandas as pd
import docx
from PyPDF2 import PdfReader
from models.student_model import Student
from utils.error_handler import handle_error, ErrorType, ErrorInfo


class FileProcessingError(Exception):
    """파일 처리 관련 예외 클래스"""
    pass


class FileService:
    """파일 처리 서비스 클래스"""
    
    # 지원하는 파일 확장자
    SUPPORTED_EXCEL_EXTENSIONS = ['.xlsx', '.xls']
    SUPPORTED_DOC_EXTENSIONS = ['.pdf', '.docx']
    SUPPORTED_IMAGE_EXTENSIONS = ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']
    
    # Excel 파일 필수 컬럼 정의
    DESCRIPTIVE_REQUIRED_COLUMNS = ['학생 이름', '반', '답안']
    MAP_REQUIRED_COLUMNS = ['학생 이름', '반']
    
    # English column name mappings
    COLUMN_MAPPINGS = {
        '학생 이름': ['학생 이름', '이름', 'Student', 'student', 'Student Name', 'student name'],
        '반': ['반', 'Class', 'class'],
        '답안': ['답안', 'Answer', 'answer']
    }
    
    def __init__(self):
        """FileService 초기화"""
        pass

    def _map_column_names(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Map English column names to Korean equivalents.
        
        Args:
            df: DataFrame with potential English column names
            
        Returns:
            DataFrame with Korean column names
        """
        print(f"DEBUG: Original columns before mapping: {list(df.columns)}")
        
        # Debug: Check each column for hidden characters
        for i, col in enumerate(df.columns):
            print(f"DEBUG: Column {i}: {repr(col)} (length: {len(col)})")
            print(f"DEBUG: Column {i} hex: {col.encode('utf-8').hex()}")
        
        # Create a copy to avoid modifying the original DataFrame
        df_mapped = df.copy()
        
        # Clean column names - remove leading/trailing whitespace and normalize
        cleaned_columns = {}
        for col in df.columns:
            cleaned_col = str(col).strip()  # Remove leading/trailing whitespace
            cleaned_columns[col] = cleaned_col
            if col != cleaned_col:
                print(f"DEBUG: Cleaned column '{col}' -> '{cleaned_col}'")
        
        # Apply column cleaning
        df_mapped = df_mapped.rename(columns=cleaned_columns)
        print(f"DEBUG: Columns after cleaning: {list(df_mapped.columns)}")
        
        # Map column names
        column_mapping = {}
        for korean_name, english_variants in self.COLUMN_MAPPINGS.items():
            for col in df_mapped.columns:
                if col in english_variants:
                    column_mapping[col] = korean_name
                    print(f"DEBUG: Mapping '{col}' -> '{korean_name}'")
                    break
        
        print(f"DEBUG: Column mapping applied: {column_mapping}")
        
        # Apply the mapping
        df_mapped = df_mapped.rename(columns=column_mapping)
        
        print(f"DEBUG: Final columns after mapping: {list(df_mapped.columns)}")
        return df_mapped

    def validate_excel_format(self, file_path: str, grading_type: str) -> Dict[str, Any]:
        """Excel 파일 형식 검증"""
        print(f"DEBUG: validate_excel_format called with grading_type='{grading_type}'")
        
        try:
            if not os.path.exists(file_path):
                error_info = handle_error(
                    FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}"),
                    ErrorType.FILE_PROCESSING,
                    context=f"validate_excel_format: {file_path}",
                    user_context="Excel 파일 검증"
                )
                return {
                    'success': False,
                    'message': error_info.user_message,
                    'data': None,
                    'error_info': error_info
                }
            
            file_extension = Path(file_path).suffix.lower()
            if file_extension not in self.SUPPORTED_EXCEL_EXTENSIONS:
                error_info = handle_error(
                    ValueError(f"지원하지 않는 파일 형식: {file_extension}"),
                    ErrorType.FILE_PROCESSING,
                    context=f"validate_excel_format: unsupported extension {file_extension}",
                    user_context="Excel 파일 형식 검증"
                )
                return {
                    'success': False,
                    'message': error_info.user_message,
                    'data': None,
                    'error_info': error_info
                }
            
            df = pd.read_excel(file_path)
            # Map column names to Korean equivalents
            df_mapped = self._map_column_names(df)
            print(f"DEBUG: Excel file loaded successfully, mapped columns: {list(df_mapped.columns)}")
            
            if df_mapped.empty:
                error_info = handle_error(
                    ValueError("Excel 파일이 비어있습니다"),
                    ErrorType.VALIDATION,
                    context=f"validate_excel_format: empty dataframe for {file_path}",
                    user_context="Excel 파일 내용 검증"
                )
                return {
                    'success': False,
                    'message': error_info.user_message,
                    'data': None,
                    'error_info': error_info
                }
            
            # Define required columns directly to avoid any encoding issues
            if grading_type == 'descriptive':
                required_columns = self.DESCRIPTIVE_REQUIRED_COLUMNS
            else:
                # Map grading only requires student name and class
                required_columns = self.MAP_REQUIRED_COLUMNS
            
            print(f"DEBUG: Required columns for '{grading_type}': {required_columns}")
            print(f"DEBUG: File columns: {list(df_mapped.columns)}")
            
            # Check if all required columns are present
            missing_columns = [col for col in required_columns if col not in df_mapped.columns]
            print(f"DEBUG: Missing columns check: {missing_columns}")
            print(f"DEBUG: Column comparison results:")
            for req_col in required_columns:
                is_present = req_col in df_mapped.columns
                print(f"  '{req_col}' in file columns: {is_present}")
            
            if missing_columns:
                print(f"DEBUG: Missing columns detected: {missing_columns}")
                print(f"DEBUG: Missing columns repr: {repr(missing_columns)}")
                error_info = handle_error(
                    ValueError(f"필수 컬럼 누락: {missing_columns}"),
                    ErrorType.VALIDATION,
                    context=f"validate_excel_format: missing columns {missing_columns}",
                    user_context="Excel 파일 컬럼 검증"
                )
                return {
                    'success': False,
                    'message': error_info.user_message,
                    'data': None,
                    'error_info': error_info
                }
            
            validation_result = self._validate_excel_data(df_mapped, grading_type)
            if not validation_result['success']:
                return validation_result
            
            return {
                'success': True,
                'message': 'Excel 파일 형식이 올바릅니다.',
                'data': df_mapped
            }

        except Exception as e:
            error_info = handle_error(
                e,
                ErrorType.FILE_PROCESSING,
                context=f"validate_excel_format: unexpected error for {file_path}",
                user_context="Excel 파일 검증"
            )
            return {
                'success': False,
                'message': error_info.user_message,
                'data': None,
                'error_info': error_info
            }

    def _validate_excel_data(self, df: pd.DataFrame, grading_type: str) -> Dict[str, Any]:
        """Excel 데이터 내용 유효성 검증"""
        try:
            # Define required columns directly to avoid any encoding issues
            if grading_type == 'descriptive':
                required_columns = self.DESCRIPTIVE_REQUIRED_COLUMNS
            else:
                # Map grading only requires student name and class
                required_columns = self.MAP_REQUIRED_COLUMNS
            
            # Check for null values in required columns
            for col in required_columns:
                if df[col].isnull().any():
                    null_rows = df[df[col].isnull()].index.tolist()
                    error_info = handle_error(
                        ValueError(f'"{col}" 컬럼에 빈 값이 있습니다. 행 번호: {null_rows}'),
                        ErrorType.VALIDATION,
                        context=f"_validate_excel_data: null values in column {col}",
                        user_context="Excel 데이터 검증"
                    )
                    return {
                        'success': False,
                        'message': error_info.user_message,
                        'data': None,
                        'error_info': error_info
                    }
            
            # Check for duplicate student names
            duplicate_names = df[df['학생 이름'].duplicated()]['학생 이름'].tolist()
            if duplicate_names:
                error_info = handle_error(
                    ValueError(f'중복된 학생 이름: {duplicate_names}'),
                    ErrorType.VALIDATION,
                    context=f"_validate_excel_data: duplicate names {duplicate_names}",
                    user_context="학생 이름 중복 검사"
                )
                return {
                    'success': False,
                    'message': error_info.user_message,
                    'data': None,
                    'error_info': error_info
                }
            
            # Check answer length for descriptive type
            if grading_type == 'descriptive':
                short_answers = []
                for idx, answer in enumerate(df['답안']):
                    if len(str(answer).strip()) < 5:
                        short_answers.append(idx + 1)
                
                if short_answers:
                    error_info = handle_error(
                        ValueError(f'답안이 너무 짧습니다 (최소 5자). 행 번호: {short_answers}'),
                        ErrorType.VALIDATION,
                        context=f"_validate_excel_data: short answers at rows {short_answers}",
                        user_context="답안 길이 검증"
                    )
                    return {
                        'success': False,
                        'message': error_info.user_message,
                        'data': None,
                        'error_info': error_info
                    }
            
            return {
                'success': True,
                'message': '데이터 유효성 검증 완료',
                'data': df
            }
            
        except Exception as e:
            error_info = handle_error(
                e,
                ErrorType.VALIDATION,
                context="_validate_excel_data: unexpected validation error",
                user_context="Excel 데이터 검증"
            )
            return {
                'success': False,
                'message': error_info.user_message,
                'data': None,
                'error_info': error_info
            }

    def process_student_data(self, excel_file_path: str, grading_type: str, 
                           image_files: Optional[List[str]] = None) -> Dict[str, Any]:
        """학생 데이터 처리 및 Student 객체 생성"""
        try:
            validation_result = self.validate_excel_format(excel_file_path, grading_type)
            if not validation_result['success']:
                return validation_result
            
            df = validation_result['data']
            students = []
            
            image_mapping = {}
            if grading_type == 'map' and image_files:
                match_result = self.match_images_to_students(df['학생 이름'].tolist(), image_files)
                if not match_result['success']:
                    return match_result
                image_mapping = match_result['mapping']
            
            for idx, row in df.iterrows():
                try:
                    student_name = str(row['학생 이름']).strip()
                    class_number = str(row['반']).strip()
                    
                    if grading_type == 'descriptive':
                        answer = str(row['답안']).strip()
                        image_path = None
                    else:
                        answer = ""
                        image_path = image_mapping.get(student_name)
                        
                        if not image_path:
                            return {
                                'success': False,
                                'message': f'학생 "{student_name}"의 이미지 파일을 찾을 수 없습니다.',
                                'students': []
                            }
                    
                    student = Student(
                        name=student_name,
                        class_number=class_number,
                        answer=answer,
                        image_path=image_path
                    )
                    students.append(student)
                    
                except Exception as e:
                    return {
                        'success': False,
                        'message': f'학생 데이터 처리 중 오류 (행 {idx + 1}): {str(e)}',
                        'students': []
                    }
            
            return {
                'success': True,
                'message': f'{len(students)}명의 학생 데이터를 성공적으로 처리했습니다.',
                'students': students
            }
            
        except Exception as e:
            raise FileProcessingError(f"학생 데이터 처리 중 오류 발생: {str(e)}")

    def match_images_to_students(self, student_names: List[str], image_files: List[str]) -> Dict[str, Any]:
        """이미지 파일과 학생 이름 매칭"""
        try:
            mapping = {}
            unmatched_students = []
            
            for student_name in student_names:
                matched = False
                student_name_clean = self._clean_name_for_matching(student_name)
                
                for image_path in image_files:
                    image_filename = Path(image_path).stem
                    image_name_clean = self._clean_name_for_matching(image_filename)
                    
                    if self._is_name_match(student_name_clean, image_name_clean):
                        mapping[student_name] = image_path
                        matched = True
                        break
                
                if not matched:
                    unmatched_students.append(student_name)
            
            if unmatched_students:
                return {
                    'success': False,
                    'message': f'다음 학생들의 이미지 파일을 찾을 수 없습니다: {", ".join(unmatched_students)}',
                    'mapping': {}
                }
            
            return {
                'success': True,
                'message': f'{len(mapping)}개의 이미지가 성공적으로 매칭되었습니다.',
                'mapping': mapping
            }
            
        except Exception as e:
            raise FileProcessingError(f"이미지 매칭 중 오류 발생: {str(e)}")

    def _clean_name_for_matching(self, name: str) -> str:
        """매칭을 위한 이름 정규화"""
        cleaned = re.sub(r'[^a-zA-Z0-9가-힣]', '', name.lower().strip())
        return cleaned

    def _is_name_match(self, student_name: str, image_name: str) -> bool:
        """학생 이름과 이미지 파일명 매칭 여부 확인"""
        if student_name == image_name:
            return True
        
        if student_name in image_name:
            return True
        
        if len(student_name) >= 2 and image_name in student_name:
            return True
        
        if len(student_name) >= 3:
            given_name = student_name[1:]
            if given_name in image_name or image_name in given_name:
                return True
        
        return False

    def extract_document_content(self, file_path: str) -> Dict[str, Any]:
        """PDF/DOCX 문서 내용 추출"""
        try:
            if not os.path.exists(file_path):
                return {
                    'success': False,
                    'message': f'파일을 찾을 수 없습니다: {file_path}',
                    'content': ''
                }
            
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension not in self.SUPPORTED_DOC_EXTENSIONS:
                return {
                    'success': False,
                    'message': f'지원하지 않는 문서 형식입니다. 지원 형식: {", ".join(self.SUPPORTED_DOC_EXTENSIONS)}',
                    'content': ''
                }
            
            content = ""
            
            if file_extension == '.pdf':
                content = self._extract_pdf_content(file_path)
            elif file_extension == '.docx':
                content = self._extract_docx_content(file_path)
            
            if not content.strip():
                return {
                    'success': False,
                    'message': '문서에서 텍스트 내용을 추출할 수 없습니다.',
                    'content': ''
                }
            
            return {
                'success': True,
                'message': f'문서 내용을 성공적으로 추출했습니다. (길이: {len(content)}자)',
                'content': content
            }
            
        except Exception as e:
            raise FileProcessingError(f"문서 내용 추출 중 오류 발생: {str(e)}")

    def _extract_pdf_content(self, file_path: str) -> str:
        """PDF 파일에서 텍스트 내용 추출"""
        try:
            content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            content += f"\n--- 페이지 {page_num + 1} ---\n"
                            content += page_text
                    except Exception as e:
                        print(f"페이지 {page_num + 1} 추출 중 오류: {str(e)}")
                        continue
            
            return content.strip()
            
        except Exception as e:
            raise FileProcessingError(f"PDF 내용 추출 중 오류: {str(e)}")

    def _extract_docx_content(self, file_path: str) -> str:
        """DOCX 파일에서 텍스트 내용 추출"""
        try:
            doc = docx.Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content += " | ".join(row_text) + "\n"
            
            return content.strip()
            
        except Exception as e:
            raise FileProcessingError(f"DOCX 내용 추출 중 오류: {str(e)}")

    def validate_image_files(self, image_files: List[str]) -> Dict[str, Any]:
        """이미지 파일들의 유효성 검증"""
        try:
            valid_files = []
            invalid_files = []
            
            for file_path in image_files:
                if not os.path.exists(file_path):
                    invalid_files.append(f"{file_path} (파일 없음)")
                    continue
                
                file_extension = Path(file_path).suffix.lower()
                if file_extension not in self.SUPPORTED_IMAGE_EXTENSIONS:
                    invalid_files.append(f"{file_path} (지원하지 않는 형식)")
                    continue
                
                file_size = os.path.getsize(file_path)
                if file_size > 50 * 1024 * 1024:  # 50MB
                    invalid_files.append(f"{file_path} (파일 크기 초과: {file_size // (1024*1024)}MB)")
                    continue
                
                valid_files.append(file_path)
            
            if invalid_files:
                return {
                    'success': False,
                    'message': f'유효하지 않은 이미지 파일들: {", ".join(invalid_files)}',
                    'valid_files': valid_files
                }
            
            return {
                'success': True,
                'message': f'{len(valid_files)}개의 이미지 파일이 유효합니다.',
                'valid_files': valid_files
            }
            
        except Exception as e:
            raise FileProcessingError(f"이미지 파일 검증 중 오류 발생: {str(e)}")

    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """파일 정보 조회"""
        try:
            if not os.path.exists(file_path):
                return {
                    'exists': False,
                    'message': '파일이 존재하지 않습니다.'
                }
            
            file_stat = os.stat(file_path)
            file_path_obj = Path(file_path)
            
            return {
                'exists': True,
                'name': file_path_obj.name,
                'extension': file_path_obj.suffix,
                'size_bytes': file_stat.st_size,
                'size_mb': round(file_stat.st_size / (1024 * 1024), 2),
                'message': '파일 정보 조회 완료'
            }
            
        except Exception as e:
            return {
                'exists': False,
                'message': f'파일 정보 조회 중 오류: {str(e)}'
            }